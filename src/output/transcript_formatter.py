"""
Transcript formatting module for generating .txt and .docx files.
"""

from typing import List, Dict
from pathlib import Path


class TranscriptFormatter:
    """Formats transcripts into .txt and .docx formats with speaker attribution."""
    
    def __init__(self):
        """Initialize transcript formatter."""
        pass
    
    def format_with_speakers(self, transcript_segments: List[Dict]) -> str:
        """
        Format transcript with speaker attribution as text.
        
        Args:
            transcript_segments: List of transcript segments with text, timestamps, speaker
            
        Returns:
            Formatted transcript text
        """
        lines = []
        for seg in transcript_segments:
            text = seg.get("text", "").strip()
            speaker_id = seg.get("speaker_id")
            # Support both 'start_time' and 'start' for timestamp
            start_time = seg.get("start_time", seg.get("start", 0.0))
            
            if not text:
                continue
            
            # Format: [Speaker] [Time] Text
            parts = []
            if speaker_id:
                parts.append(f"[{speaker_id}]")
            parts.append(f"[{self._format_time(start_time)}]")
            parts.append(text)
            
            lines.append(" ".join(parts))
        
        return "\n".join(lines)
    
    def format_to_txt(self, 
                     transcript_segments: List[Dict], 
                     output_path: str,
                     include_timestamps: bool = True,
                     include_speakers: bool = True) -> str:
        """
        Format transcript to plain text file.
        
        Args:
            transcript_segments: List of transcript segments with text, timestamps, speaker
            output_path: Path to save TXT file
            include_timestamps: Whether to include timestamps
            include_speakers: Whether to include speaker attribution
            
        Returns:
            Path to saved TXT file
        """
        lines = []
        for seg in transcript_segments:
            text = seg.get("text", "").strip()
            if not text:
                continue
            
            parts = []
            if include_speakers and seg.get("speaker_id"):
                parts.append(f"[{seg['speaker_id']}]")
            if include_timestamps:
                start_time = seg.get("start_time", 0.0)
                parts.append(f"[{self._format_time(start_time)}]")
            parts.append(text)
            
            lines.append(" ".join(parts))
        
        content = "\n".join(lines)
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
        output_path_obj.write_text(content, encoding='utf-8')
        return str(output_path_obj)
    
    def _format_time(self, seconds: float) -> str:
        """Format seconds to HH:MM:SS."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    
    def format_to_docx(self, 
                      transcript_segments: List[Dict], 
                      output_path: str,
                      include_timestamps: bool = True,
                      include_speakers: bool = True) -> str:
        """
        Format transcript to DOCX file.
        
        Args:
            transcript_segments: List of transcript segments
            output_path: Path to save DOCX file
            include_timestamps: Whether to include timestamps
            include_speakers: Whether to include speaker attribution
            
        Returns:
            Path to saved DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install it with: pip install python-docx")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Video Transcript', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata paragraph
        doc.add_paragraph(f'Generated transcript with {len(transcript_segments)} segments')
        doc.add_paragraph('')  # Empty line
        
        # Process segments
        for seg in transcript_segments:
            text = seg.get("text", "").strip()
            if not text:
                continue
            
            # Create paragraph for this segment
            para = doc.add_paragraph()
            
            # Add speaker if available
            if include_speakers and seg.get("speaker_id"):
                speaker_run = para.add_run(f"[{seg['speaker_id']}] ")
                speaker_run.bold = True
                speaker_run.font.color.rgb = RGBColor(0, 100, 200)  # Blue color
            
            # Add timestamp if available
            if include_timestamps:
                start_time = seg.get("start_time", seg.get("start", 0.0))
                time_run = para.add_run(f"[{self._format_time(start_time)}] ")
                time_run.italic = True
                time_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
            
            # Add transcript text
            text_run = para.add_run(text)
            text_run.font.size = Pt(11)
            
            # Add spacing between segments
            para.space_after = Pt(6)
        
        # Save document
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path_obj))
        
        return str(output_path_obj)
    
    def format_summary_to_txt(self, summary: str, output_path: str) -> str:
        """
        Format summary to text file.
        
        Args:
            summary: Summary text
            output_path: Path to save summary file
            
        Returns:
            Path to saved file
        """
        # TODO: Implement summary formatting
        pass
    
    def format_summary_to_pdf(self, summary: str, output_path: str, key_points: List[str] = None) -> str:
        """
        Format summary to PDF file.
        
        Args:
            summary: Summary text
            output_path: Path to save PDF file
            key_points: Optional list of key points to include
            
        Returns:
            Path to saved PDF file
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            from reportlab.lib.colors import HexColor
        except ImportError:
            raise ImportError("reportlab is required for PDF export. Install it with: pip install reportlab")
        
        # Create PDF document
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
        
        doc = SimpleDocTemplate(
            str(output_path_obj),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Heading style
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=20
        )
        
        # Body style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            textColor=HexColor('#333333'),
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leading=16
        )
        
        # Key point style
        key_point_style = ParagraphStyle(
            'CustomKeyPoint',
            parent=styles['Normal'],
            fontSize=11,
            textColor=HexColor('#333333'),
            leftIndent=20,
            spaceAfter=8,
            leading=16
        )
        
        # Add title
        title = Paragraph("Video Summary", title_style)
        elements.append(title)
        elements.append(Spacer(1, 0.2*inch))
        
        # Add key points if available
        if key_points and len(key_points) > 0:
            heading = Paragraph("Key Points", heading_style)
            elements.append(heading)
            
            for point in key_points:
                # Clean up key point text
                point_text = point.strip()
                if point_text:
                    # Add bullet point
                    bullet_point = Paragraph(f"• {point_text}", key_point_style)
                    elements.append(bullet_point)
            
            elements.append(Spacer(1, 0.3*inch))
        
        # Add summary heading
        summary_heading = Paragraph("Full Summary", heading_style)
        elements.append(summary_heading)
        
        # Split summary into paragraphs and add them
        summary_paragraphs = summary.split('\n\n')
        for para_text in summary_paragraphs:
            if para_text.strip():
                # Clean and escape HTML special characters
                para_text_clean = para_text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                para = Paragraph(para_text_clean, body_style)
                elements.append(para)
                elements.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(elements)
        
        return str(output_path_obj)
    
    def create_transcript_with_visual_references(self, 
                                                transcript_segments: List[Dict],
                                                visual_cues: List[Dict]) -> List[Dict]:
        """
        Create transcript with visual references (slides, gestures, etc.).
        
        Args:
            transcript_segments: List of transcript segments
            visual_cues: List of visual cues with timestamps
            
        Returns:
            Enhanced transcript segments with visual references
        """
        # TODO: Implement visual reference integration
        pass

